# gemini_app/gemini.py
import google.generativeai as genai
from PIL import Image
import os
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from io import BytesIO
from dotenv import load_dotenv

load_dotenv()

class geminiApi:
    def __init__(self):
        self.api_key = os.getenv("GEMINI_API_KEY")
        self.model = None

    def configure(self):
        genai.configure(api_key=self.api_key)
        self.model = genai.GenerativeModel(model_name="gemini-pro-vision")
        return self.model

    def configure_2(self):
        genai.configure(api_key=self.api_key)
        self.model = genai.GenerativeModel(model_name="gemini-pro")
        return self.model
    
    def generate_user_story(self, file_path, description, file_type):
        if file_type == 'image':
            imagen = Image.open(file_path)
            prompt = f"{description}, quiero que formules dos historias de usuario con el formato de como, quiero, para (no debe especificar la sintaxis)"
            response = self.model.generate_content([prompt, imagen])
        else:
            return []

        historias_usuario = response.text.split('\n')
        return [historia.strip() for historia in historias_usuario if historia.strip()]

def generate_word_document(historias_usuario):
    doc = Document()
    doc.add_heading('Historias de Usuario', 0)
    for historia in historias_usuario:
        doc.add_paragraph(historia)
    byte_io = BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)
    return byte_io.getvalue()

def generate_pdf_document(historias_usuario):
    byte_io = BytesIO()
    c = canvas.Canvas(byte_io, pagesize=letter)
    width, height = letter
    c.drawString(100, height - 40, "Historias de Usuario")
    y = height - 60
    for historia in historias_usuario:
        c.drawString(100, y, historia)
        y -= 20
        if y < 40:
            c.showPage()
            y = height - 40
    c.save()
    byte_io.seek(0)
    return byte_io.getvalue()
